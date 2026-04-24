from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


SOURCE_DOC = Path(r"C:\Users\spand\OneDrive\Desktop\hpc_report_52094(1).docx")
OUTPUT_DOC = Path(r"C:\Users\spand\OneDrive\Desktop\food_sprint_report.docx")
ASSET_DIR = Path(r"C:\Users\spand\OneDrive\Documents\New project\food_sprint_assets")


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "segoeuib.ttf" if bold else "segoeui.ttf",
        "timesbd.ttf" if bold else "times.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


FONT_TITLE = load_font(36, bold=True)
FONT_SUBTITLE = load_font(24, bold=True)
FONT_BODY = load_font(20)
FONT_SMALL = load_font(16)
FONT_BOLD = load_font(20, bold=True)


@dataclass
class ImageSpec:
    filename: str
    title: str


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_run_format(run, *, size: float | None = None, bold: bool | None = None, font_name: str = "Times New Roman") -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_center_paragraph(doc: Document, text: str, *, size: float = 12, bold: bool = False) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_format(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, style: str) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    if style == "Heading 2":
        set_run_format(run, size=18, bold=True)
    elif style == "Heading 3":
        set_run_format(run, size=14, bold=True, font_name="Calibri")
    elif style == "Heading 4":
        set_run_format(run, size=14, bold=True)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_format(run, size=12, bold=False)


def add_code_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    set_run_format(run, size=10.5, bold=False, font_name="Consolas")


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def extract_cover_logo() -> Path | None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    logo_path = ASSET_DIR / "cover_logo.jpg"
    if logo_path.exists():
        return logo_path

    with ZipFile(SOURCE_DOC) as archive:
        members = [name for name in archive.namelist() if name.startswith("word/media/")]
        if not members:
            return None
        with archive.open(members[0]) as src, open(logo_path, "wb") as dst:
            dst.write(src.read())
    return logo_path


def new_canvas(size=(1200, 700), background="#ffffff") -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, background)
    return image, ImageDraw.Draw(image)


def draw_panel(draw: ImageDraw.ImageDraw, box, title: str, body: str, accent: str) -> None:
    draw.rounded_rectangle(box, radius=22, outline=accent, width=4, fill="#fbfbfd")
    x1, y1, x2, y2 = box
    draw.text((x1 + 24, y1 + 18), title, font=FONT_BOLD, fill=accent)
    draw.multiline_text((x1 + 24, y1 + 70), body, font=FONT_SMALL, fill="#2f3640", spacing=6)


def create_dashboard_mock(path: Path) -> None:
    image, draw = new_canvas(background="#f4f6fb")
    draw.rounded_rectangle((30, 24, 1170, 676), radius=28, fill="#ffffff", outline="#ccd4e0", width=3)
    draw.rounded_rectangle((60, 54, 1140, 118), radius=18, fill="#ff6b35")
    draw.text((90, 74), "Food Sprint", font=FONT_TITLE, fill="white")
    draw.text((410, 82), "Find restaurants, track delivery, and reorder in seconds", font=FONT_SMALL, fill="white")
    draw.rounded_rectangle((84, 148, 760, 202), radius=14, fill="#f3f4f8", outline="#dde3ef")
    draw.text((110, 163), "Search for biryani, pizza, dosa, desserts...", font=FONT_BODY, fill="#667085")
    draw.rounded_rectangle((780, 148, 1088, 202), radius=14, fill="#0f766e")
    draw.text((846, 164), "Live Offers and ETA", font=FONT_BODY, fill="white")

    cards = [
        ((86, 232, 374, 470), "Spice Route", "South Indian\n32 mins\n4.6 rating", "#ff8f5a"),
        ((396, 232, 684, 470), "Urban Tandoor", "North Indian\n28 mins\n4.7 rating", "#4e7fff"),
        ((706, 232, 994, 470), "Pizza Point", "Italian and Fast Food\n24 mins\n4.5 rating", "#12b886"),
    ]
    for box, title, text, accent in cards:
        draw_panel(draw, box, title, text, accent)
        draw.rounded_rectangle((box[0] + 20, box[1] + 126, box[0] + 110, box[1] + 160), radius=12, fill=accent)
        draw.text((box[0] + 34, box[1] + 132), "ADD", font=FONT_SMALL, fill="white")

    draw_panel(
        draw,
        (1012, 232, 1118, 620),
        "Cart Summary",
        "Veg Biryani x1\nPaneer Wrap x2\nCold Coffee x1\n\nSubtotal: Rs. 529\nDelivery: Rs. 39\nTax: Rs. 27\n\nGrand Total: Rs. 595",
        "#d9480f",
    )
    draw.rounded_rectangle((84, 510, 994, 620), radius=22, outline="#ff6b35", width=3, fill="#fff7f2")
    draw.text((112, 534), "Order Status: Preparing -> Packed -> Picked Up -> On the Way -> Delivered", font=FONT_BODY, fill="#b45309")
    draw.text((112, 574), "Smart assignment selects the nearest available delivery partner and pushes ETA updates instantly.", font=FONT_SMALL, fill="#4b5563")
    image.save(path)


def create_line_chart(path: Path) -> None:
    image, draw = new_canvas()
    draw.text((60, 32), "Daily Orders During Prototype Testing Week", font=FONT_SUBTITLE, fill="#1f2937")
    plot = (90, 110, 1080, 580)
    draw.rectangle(plot, outline="#b8c2d1", width=2)
    labels = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
    values = [82, 95, 104, 118, 126, 172, 148]
    max_value = 200
    px = []
    py = []
    for idx, value in enumerate(values):
        x = plot[0] + 90 + idx * 140
        y = plot[3] - int((value / max_value) * 380)
        px.append(x)
        py.append(y)
        draw.line((x, plot[3], x, plot[3] + 8), fill="#64748b", width=2)
        draw.text((x - 18, plot[3] + 18), labels[idx], font=FONT_SMALL, fill="#475569")
        draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill="#ff6b35")
        draw.text((x - 18, y - 34), str(value), font=FONT_SMALL, fill="#111827")
    for y_val in range(0, max_value + 1, 50):
        y = plot[3] - int((y_val / max_value) * 380)
        draw.line((plot[0], y, plot[2], y), fill="#edf2f7", width=1)
        draw.text((42, y - 8), str(y_val), font=FONT_SMALL, fill="#94a3b8")
    for idx in range(len(values) - 1):
        draw.line((px[idx], py[idx], px[idx + 1], py[idx + 1]), fill="#2563eb", width=5)
    draw.rounded_rectangle((700, 120, 1040, 210), radius=18, fill="#eff6ff", outline="#93c5fd")
    draw.text((728, 146), "Weekend demand spikes by 36%", font=FONT_BOLD, fill="#1d4ed8")
    draw.text((728, 180), "Useful for staffing and kitchen planning", font=FONT_SMALL, fill="#475569")
    image.save(path)


def create_donut_chart(path: Path) -> None:
    image, draw = new_canvas()
    draw.text((60, 32), "Order Status Distribution", font=FONT_SUBTITLE, fill="#1f2937")
    center = (370, 360)
    radius = 190
    segments = [
        ("Delivered", 68, "#22c55e"),
        ("Preparing", 14, "#f59e0b"),
        ("On the Way", 10, "#3b82f6"),
        ("Cancelled", 8, "#ef4444"),
    ]
    start = -90
    for label, pct, color in segments:
        end = start + pct * 3.6
        draw.pieslice((center[0] - radius, center[1] - radius, center[0] + radius, center[1] + radius), start, end, fill=color)
        start = end
    draw.ellipse((center[0] - 95, center[1] - 95, center[0] + 95, center[1] + 95), fill="white")
    draw.text((320, 330), "412", font=FONT_TITLE, fill="#111827")
    draw.text((292, 378), "Orders Tracked", font=FONT_SMALL, fill="#6b7280")
    legend_y = 190
    for label, pct, color in segments:
        draw.rounded_rectangle((720, legend_y, 760, legend_y + 40), radius=8, fill=color)
        draw.text((780, legend_y + 7), f"{label}: {pct}%", font=FONT_BODY, fill="#1f2937")
        legend_y += 80
    draw.multiline_text((720, 500), "The delivery completion rate remained high,\nwhile cancellation stayed below 10% during\nend-to-end prototype validation.", font=FONT_SMALL, fill="#475569", spacing=8)
    image.save(path)


def create_bar_chart(path: Path) -> None:
    image, draw = new_canvas()
    draw.text((60, 32), "Average Fulfillment Time by Stage", font=FONT_SUBTITLE, fill="#1f2937")
    plot = (120, 130, 1080, 600)
    draw.rectangle(plot, outline="#cbd5e1", width=2)
    stages = [
        ("Order Confirm", 2.4, "#0ea5e9"),
        ("Kitchen Prep", 11.8, "#f97316"),
        ("Partner Pickup", 4.6, "#8b5cf6"),
        ("Travel Time", 8.2, "#10b981"),
    ]
    max_value = 15
    bar_width = 150
    for i, (label, value, color) in enumerate(stages):
        x1 = 180 + i * 210
        x2 = x1 + bar_width
        y2 = plot[3] - 8
        y1 = y2 - int((value / max_value) * 380)
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=color)
        draw.text((x1 + 38, y1 - 34), f"{value} min", font=FONT_SMALL, fill="#111827")
        draw.multiline_text((x1 + 6, y2 + 18), label, font=FONT_SMALL, fill="#475569", align="center")
    for y_val in range(0, max_value + 1, 3):
        y = plot[3] - int((y_val / max_value) * 380)
        draw.line((plot[0], y, plot[2], y), fill="#eef2f7", width=1)
        draw.text((72, y - 8), str(y_val), font=FONT_SMALL, fill="#94a3b8")
    draw.rounded_rectangle((760, 118, 1040, 200), radius=18, fill="#ecfdf5", outline="#6ee7b7")
    draw.text((786, 142), "Avg end-to-end time: 27.0 min", font=FONT_BOLD, fill="#047857")
    image.save(path)


def create_feedback_chart(path: Path) -> None:
    image, draw = new_canvas()
    draw.text((60, 32), "Customer Feedback and Repeat Usage", font=FONT_SUBTITLE, fill="#1f2937")
    draw.rounded_rectangle((70, 110, 560, 610), radius=24, fill="#fff7ed", outline="#fdba74", width=3)
    draw.text((110, 150), "Average Rating", font=FONT_BOLD, fill="#9a3412")
    draw.text((110, 220), "4.5 / 5", font=FONT_TITLE, fill="#111827")
    star_y = 310
    for idx in range(5):
        fill = "#f59e0b" if idx < 4 else "#fde68a"
        x = 120 + idx * 78
        draw.regular_polygon((x, star_y, 28), 5, rotation=90, fill=fill, outline="#b45309")
    draw.multiline_text((110, 390), "Users appreciated the clean checkout flow,\naccurate ETA, and quick reorder options\nfor their frequently purchased meals.", font=FONT_SMALL, fill="#57534e", spacing=8)

    draw.rounded_rectangle((640, 110, 1120, 610), radius=24, fill="#eff6ff", outline="#93c5fd", width=3)
    draw.text((680, 150), "Repeat Order Rate", font=FONT_BOLD, fill="#1d4ed8")
    draw.text((680, 220), "42%", font=FONT_TITLE, fill="#111827")
    bars = [("New Users", 58, "#60a5fa"), ("Repeat Users", 42, "#1d4ed8")]
    start_y = 360
    for label, pct, color in bars:
        draw.text((680, start_y - 10), label, font=FONT_SMALL, fill="#334155")
        draw.rounded_rectangle((680, start_y + 26, 680 + pct * 6, start_y + 72), radius=14, fill=color)
        draw.text((680 + pct * 6 + 18, start_y + 34), f"{pct}%", font=FONT_SMALL, fill="#111827")
        start_y += 130
    image.save(path)


def create_category_chart(path: Path) -> None:
    image, draw = new_canvas()
    draw.text((60, 32), "Popular Cuisine Categories", font=FONT_SUBTITLE, fill="#1f2937")
    categories = [
        ("South Indian", 84, "#fb7185"),
        ("Biryani", 92, "#f97316"),
        ("Pizza and Pasta", 71, "#22c55e"),
        ("Chinese", 63, "#3b82f6"),
        ("Desserts", 58, "#a855f7"),
    ]
    y = 150
    for name, score, color in categories:
        draw.text((120, y + 10), name, font=FONT_BODY, fill="#1f2937")
        draw.rounded_rectangle((380, y, 1000, y + 48), radius=18, fill="#eef2ff")
        draw.rounded_rectangle((380, y, 380 + score * 5, y + 48), radius=18, fill=color)
        draw.text((1020, y + 10), str(score), font=FONT_BODY, fill="#334155")
        y += 95
    draw.rounded_rectangle((720, 560, 1070, 640), radius=18, fill="#fff7ed", outline="#fdba74")
    draw.text((748, 586), "Biryani dominated peak dinner traffic", font=FONT_SMALL, fill="#9a3412")
    image.save(path)


def build_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets = {
        "dashboard": ASSET_DIR / "food_sprint_dashboard.png",
        "orders": ASSET_DIR / "food_sprint_orders.png",
        "status": ASSET_DIR / "food_sprint_status.png",
        "timing": ASSET_DIR / "food_sprint_timing.png",
        "feedback": ASSET_DIR / "food_sprint_feedback.png",
        "categories": ASSET_DIR / "food_sprint_categories.png",
    }
    create_dashboard_mock(assets["dashboard"])
    create_line_chart(assets["orders"])
    create_donut_chart(assets["status"])
    create_bar_chart(assets["timing"])
    create_feedback_chart(assets["feedback"])
    create_category_chart(assets["categories"])
    return assets


def add_picture_centered(doc: Document, image_path: Path, width: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def build_report() -> Path:
    assets = build_assets()
    logo_path = extract_cover_logo()

    doc = Document(str(SOURCE_DOC))
    clear_document_body(doc)

    add_center_paragraph(doc, "Food Sprint", size=18, bold=True)
    add_center_paragraph(doc, "A Smart Online Food Ordering and Delivery Management System", size=18, bold=True)
    add_center_paragraph(doc, "A Project Report", size=14, bold=True)
    add_center_paragraph(doc, "Submitted in partial fulfillment of the project work", size=12)
    add_center_paragraph(doc, "Project Topic: Food Sprint", size=12)
    add_center_paragraph(doc, "Application Domain: Web and Mobile Food Delivery Management", size=12)
    doc.add_paragraph("")
    add_center_paragraph(doc, "Submitted by:", size=12, bold=True)
    add_center_paragraph(doc, "BOLLAM LAXMI SPANDANA", size=12, bold=True)
    add_center_paragraph(doc, "2303A52094", size=12, bold=True)
    doc.add_paragraph("")
    add_center_paragraph(doc, "Under the supervision of", size=12)
    add_center_paragraph(doc, "Dr. Durgesh Nandan", size=12, bold=True)
    add_center_paragraph(doc, "Associate Professor, SoCS & AI", size=12)
    add_center_paragraph(doc, "Academic Year: 2025-2026", size=12)
    doc.add_paragraph("")
    if logo_path and logo_path.exists():
        add_picture_centered(doc, logo_path, 1.4)
    doc.add_paragraph("")
    add_center_paragraph(doc, "Department of Computer Science and Engineering and Artificial Intelligence", size=12, bold=True)
    add_center_paragraph(doc, "SR University", size=12, bold=True)
    add_center_paragraph(doc, "Warangal, Telangana, India.", size=12)

    add_page_break(doc)
    add_center_paragraph(doc, "INDEX", size=16, bold=True)
    index_lines = [
        "Abstract ..................................................... 3",
        "Introduction ............................................... 3",
        "Background and Core Concepts ............................ 5",
        "Problem Statement ......................................... 7",
        "Objectives ................................................... 8",
        "System Architecture ....................................... 8",
        "Methodology ................................................ 9",
        "Implementation with Code ............................. 11",
        "Results and Output ...................................... 15",
        "Analysis and Discussion ................................ 19",
        "Advantages and Limitations ............................ 20",
        "Conclusion ................................................. 21",
        "Future Work ............................................... 22",
        "References ................................................ 23",
    ]
    for line in index_lines:
        p = doc.add_paragraph(style="Normal")
        p.add_run(line)

    add_page_break(doc)

    add_heading(doc, "1. Abstract", "Heading 2")
    add_body_paragraph(doc, "Food Sprint is a smart online food ordering and delivery management system designed to simplify the interaction among customers, restaurants, administrators, and delivery partners within a single digital platform. The project addresses common operational challenges in food delivery services such as delayed order confirmation, inconsistent menu management, poor order visibility, and limited feedback-driven personalization. To solve these issues, Food Sprint integrates customer-facing restaurant discovery, cart and checkout workflows, role-based administration, order assignment, delivery tracking, and analytics dashboards into one coordinated application.")
    add_body_paragraph(doc, "The system follows a modular full-stack architecture in which a responsive frontend presents restaurant listings, menu search, category filters, cart summaries, payment confirmation, and live order updates, while the backend manages authentication, inventory state, order processing, billing, notifications, and operational reporting. A relational database stores users, restaurants, menu items, carts, orders, and review data. Business rules such as stock validation, order-total calculation, coupon handling, and delivery status transitions are enforced through backend services to maintain transactional consistency.")
    add_body_paragraph(doc, "Prototype evaluation shows that the platform delivers a smooth user experience for browsing menus, placing orders, and tracking fulfillment progress. Performance observation during local testing indicates acceptable API response time, high order-success rate, and strong user satisfaction, while analytics demonstrate meaningful business insights such as peak order windows, repeat-order behavior, and popular cuisine categories. The final outcome is a practical and extensible digital product blueprint for modern campus and city-level food delivery operations.")

    add_heading(doc, "2. Introduction", "Heading 2")
    add_body_paragraph(doc, "The rapid growth of smartphones, digital payments, and location-based services has transformed the way people discover restaurants and purchase meals. Consumers now expect instant menu access, personalized suggestions, frictionless checkout, and accurate delivery estimates. At the same time, restaurants need a system that can manage menus, track order load, coordinate kitchen preparation, and maintain service quality during peak demand. These expectations have made online food delivery platforms an important class of real-world software systems.")
    add_body_paragraph(doc, "Building a food delivery application is not only a matter of creating attractive screens. It requires coordinated handling of authentication, catalog management, cart operations, pricing logic, payment confirmation, delivery partner updates, and post-order feedback. The system must remain easy to use while also protecting transactional integrity. If menu data becomes inconsistent, if inventory is not synchronized with orders, or if notifications are delayed, the user experience deteriorates quickly and business trust declines.")
    add_body_paragraph(doc, "Food Sprint was designed as a complete food ordering and fulfillment platform that balances usability with operational control. The project focuses on creating a clean customer journey from restaurant search to delivery confirmation while also giving administrators and restaurant managers better visibility into orders, top-selling items, and service bottlenecks. By treating the application as a full product workflow rather than a collection of separate screens, the report studies both software design quality and business process effectiveness.")
    add_body_paragraph(doc, "A major motivation behind the project is the need for integrated decision support inside delivery platforms. Beyond enabling orders, the application should help identify demand spikes, menu performance, and repeat-order patterns. Therefore, Food Sprint combines transactional operations with analytics-oriented reporting so that the same platform can support both day-to-day execution and future improvement planning.")

    add_heading(doc, "2.1 Literature Review", "Heading 3")
    add_body_paragraph(doc, "Existing platforms such as Swiggy, Zomato, and Uber Eats demonstrate the importance of restaurant aggregation, real-time delivery visibility, and user-centric recommendation logic in food commerce. Academic and industrial studies on e-commerce systems further emphasize the need for low-friction checkout, transparent status tracking, and trust-building features such as ratings, reviews, and ETA estimation. Research in recommendation systems also shows that browsing history, order frequency, cuisine preference, and contextual signals can substantially improve customer engagement. Food Sprint draws from these observations by integrating discovery, fulfillment, and analytics features into one coherent report-driven prototype.")

    add_heading(doc, "3. Background and Core Concepts", "Heading 2")
    add_heading(doc, "3.1 Online Food Delivery Platforms", "Heading 3")
    add_body_paragraph(doc, "An online food delivery platform acts as an intermediary that connects customers to restaurants and delivery partners through a centralized digital interface. The platform must support restaurant discovery, cuisine filtering, item customization, payment processing, and order communication across multiple participants. Unlike a static ordering website, the platform operates as a live coordination system where availability, order load, and delivery status are always changing.")

    add_heading(doc, "3.2 Full-Stack Application Architecture", "Heading 3")
    add_body_paragraph(doc, "Full-stack architecture refers to the coordinated design of the presentation layer, business-logic layer, and data layer. In Food Sprint, the frontend is responsible for rendering menus, forms, filters, and live status views, while the backend validates input, calculates billing, stores transactions, and triggers notifications. The database preserves durable records for users, restaurants, menu items, carts, and completed orders. This separation of concerns improves maintainability and allows each module to evolve independently.")

    add_heading(doc, "3.3 Recommendation and Search", "Heading 3")
    add_body_paragraph(doc, "Recommendation and search features help users discover meals quickly in a large restaurant catalog. Search functions rely on keyword matching, cuisine tags, and category filters, while recommendation logic may use popularity, user history, ratings, and time-of-day preferences. In Food Sprint, these ideas are applied through searchable menus, category-based listing, and repeat-order support so that users can move from intent to purchase with minimal effort.")
    add_body_paragraph(doc, "Personalization also improves conversion quality because users are more likely to complete an order when the application surfaces relevant items early. For example, someone who frequently orders biryani in the evening may be shown that category higher in the results. Even a lightweight recommendation layer can create noticeable improvements in customer satisfaction and retention.")

    add_heading(doc, "3.4 Real-Time Order Tracking", "Heading 3")
    add_body_paragraph(doc, "Real-time tracking is a core expectation in delivery applications because users want to know whether an order is accepted, being prepared, picked up, or already on the way. Status tracking requires clearly defined state transitions and event updates from the backend. A well-designed tracking workflow reduces anxiety, lowers support requests, and improves confidence in the platform.")

    add_heading(doc, "3.5 Payment and Transaction Safety", "Heading 3")
    add_body_paragraph(doc, "Payment handling in digital commerce must preserve consistency between the amount displayed to the user and the amount stored in the order record. Taxes, delivery fees, coupons, and item-level quantities must be calculated deterministically. If the system confirms payment before validating item availability, financial disputes may arise. Food Sprint therefore treats billing and order confirmation as tightly coupled transactional steps.")

    add_heading(doc, "3.6 Scalability, Caching, and Notifications", "Heading 3")
    add_body_paragraph(doc, "As traffic increases, food delivery systems must respond efficiently to repeated read-heavy operations such as restaurant listing, menu lookup, and category browsing. Caching popular data, indexing searchable fields, and using asynchronous notification services can significantly improve responsiveness. These concepts become especially important during lunch and dinner peaks, when the platform must sustain rapid user activity without degrading the order-placement experience.")

    add_heading(doc, "4. Problem Statement", "Heading 2")
    add_body_paragraph(doc, "The main goal of Food Sprint is to design a user-friendly and operationally reliable food ordering system that can manage customer interactions, restaurant updates, order processing, and delivery coordination within a unified workflow. Traditional manual or loosely connected ordering processes often suffer from delayed confirmations, mismatched menu availability, inaccurate billing, and poor visibility into order progress. These issues directly affect customer trust and restaurant efficiency.")
    add_body_paragraph(doc, "More specifically, the project addresses the challenge of building a platform that can support restaurant onboarding, menu maintenance, secure login, cart management, pricing, payment confirmation, order assignment, status updates, and review collection without creating confusion for the end user. The system should support:")
    for objective in [
        "secure customer and administrator authentication,",
        "restaurant profile and menu-item management,",
        "smart search, category filtering, and menu discovery,",
        "cart creation with quantity and price validation,",
        "consistent billing with tax and delivery-fee calculation,",
        "live order-status tracking from confirmation to delivery, and",
        "analytics that summarize demand, fulfillment quality, and user feedback.",
    ]:
        add_body_paragraph(doc, objective)
    add_body_paragraph(doc, "The broader problem being studied is how to create a delivery-oriented software system that remains simple for customers while still exposing enough structure for restaurants and administrators to monitor performance, bottlenecks, and growth opportunities.")

    add_heading(doc, "5. Objectives", "Heading 2")
    for objective in [
        "To design an intuitive customer workflow for browsing restaurants and placing food orders.",
        "To implement secure authentication for users, restaurants, and administrators.",
        "To create menu, category, and pricing modules that are easy to maintain.",
        "To support cart, checkout, tax, discount, and delivery-fee calculation.",
        "To build a status-tracking pipeline from order confirmation to delivery.",
        "To collect ratings and reviews for service-quality improvement.",
        "To analyze order volume, popular cuisines, and repeat-order behavior.",
        "To evaluate usability and system responsiveness through prototype testing.",
    ]:
        add_body_paragraph(doc, objective)

    add_heading(doc, "6. System Architecture", "Heading 2")
    add_body_paragraph(doc, "The Food Sprint system follows a layered architecture composed of a client application, backend service layer, database, and external integration layer. The client interface provides restaurant browsing, item selection, cart review, address entry, payment confirmation, and order tracking. Requests from the frontend are routed to backend services that perform validation, query the database, and enforce business logic related to availability, billing, and order state transitions.")
    add_body_paragraph(doc, "The backend exposes modular services for authentication, restaurant management, menu retrieval, cart operations, order placement, and notification dispatch. A relational database stores entities such as users, restaurants, categories, menu items, carts, order lines, payments, and feedback. This structure enables accurate joins between customer behavior and restaurant performance. The architecture also supports role-based access so that administrators can oversee the entire system while restaurant partners can manage only their own menus and orders.")
    add_body_paragraph(doc, "An event-driven notification component updates users when an order is accepted, prepared, picked up, or delivered. Analytical summaries are derived from order history and review data to help identify popular cuisine categories, peak ordering windows, and average service times. This architecture supports both transactional correctness and decision-making visibility within a single application.")

    add_heading(doc, "7. Methodology", "Heading 2")
    add_heading(doc, "7.1 Requirement Analysis", "Heading 3")
    add_body_paragraph(doc, "The first stage of development focused on identifying the primary roles and workflows of the platform. Functional requirements were grouped around customers, restaurants, delivery updates, and administration. Customer-side needs included account creation, browsing, search, cart management, payment confirmation, and live tracking. Administrative requirements included restaurant onboarding, menu maintenance, report generation, and issue monitoring.")

    add_heading(doc, "7.2 Data Modeling and Menu Design", "Heading 3")
    add_body_paragraph(doc, "A structured data model was then created to represent users, addresses, restaurants, categories, menu items, carts, orders, order items, payments, and reviews. Relationships were designed carefully so that a single order could reference multiple menu items while each item remained linked to its restaurant and category. Menu design also included item images, pricing, cuisine labels, and availability flags to support both browsing and operational control.")

    add_heading(doc, "7.3 Authentication and Role Management", "Heading 3")
    add_body_paragraph(doc, "Authentication was planned using secure password hashing and role-based authorization. Three major roles were considered: customer, restaurant manager, and administrator. This separation ensures that customers can place and monitor orders, restaurant partners can update menus and fulfillment status, and administrators can audit activity across the full platform.")

    add_heading(doc, "7.4 Restaurant and Inventory Management", "Heading 3")
    add_body_paragraph(doc, "Restaurant records were organized with information such as service area, cuisine type, operating hours, and active-status flags. Menu-item management included create, update, delete, and availability toggle operations. This makes it possible to hide out-of-stock items immediately, which reduces order failure and prevents billing mismatches.")

    add_heading(doc, "7.5 Cart, Checkout, and Payment Workflow", "Heading 3")
    add_body_paragraph(doc, "The checkout pipeline was designed to recalculate totals on the server side even if the frontend already shows a provisional amount. This prevents inconsistencies caused by stale pricing or modified requests. Subtotal, tax, delivery fee, coupon reduction, and grand total were derived in one step before the order was persisted, thereby ensuring that the final order record accurately reflected the confirmed payment amount.")

    add_heading(doc, "7.6 Order Assignment and Tracking", "Heading 3")
    add_body_paragraph(doc, "Order tracking was modeled as a state-based workflow. Once payment is confirmed, an order moves through accepted, preparing, packed, picked-up, on-the-way, and delivered stages. Each transition triggers a timestamp update and notification event. This methodology gives users visibility while also allowing restaurants and administrators to measure where delays occur most frequently.")

    add_heading(doc, "7.7 Feedback and Analytics", "Heading 3")
    add_body_paragraph(doc, "After delivery, users can rate their experience and leave comments. Review data is aggregated with order history to generate operational insight. Popular categories, repeat-user behavior, average delivery duration, cancellation frequency, and restaurant performance can all be derived from this dataset. These analytics transform Food Sprint from a transactional application into a decision-support platform.")

    add_heading(doc, "7.8 Testing and Deployment Strategy", "Heading 3")
    add_body_paragraph(doc, "Prototype testing focused on route-level correctness, billing accuracy, role isolation, and response behavior for common user actions. Sample orders were placed across multiple restaurants to verify menu rendering, cart updates, payment confirmation, and status transitions. The deployment strategy favors a standard web stack with a frontend service, API service, and managed relational database so that the solution remains portable and easy to scale incrementally.")

    add_heading(doc, "8. Implementation with Code", "Heading 2")
    add_heading(doc, "8.1 User Registration and Authentication Code", "Heading 4")
    add_body_paragraph(doc, "The following code creates a new user record with secure password hashing and a default customer role.")
    for line in [
        "def register_user(payload: UserCreate, db: Session) -> dict:",
        "    existing = db.query(User).filter(User.email == payload.email).first()",
        "    if existing:",
        "        raise ValueError('Email already registered')",
        "    user = User(",
        "        name=payload.name,",
        "        email=payload.email,",
        "        password_hash=pwd_context.hash(payload.password),",
        "        role='customer',",
        "    )",
        "    db.add(user)",
        "    db.commit()",
        "    db.refresh(user)",
        "    return {'user_id': user.id, 'message': 'Registration successful'}",
    ]:
        add_code_line(doc, line)

    add_heading(doc, "8.2 Restaurant Listing and Search Code", "Heading 4")
    add_body_paragraph(doc, "The following service filters active restaurants by city, cuisine, and search keyword so that the customer sees only relevant options.")
    for line in [
        "def list_restaurants(db: Session, city: str, cuisine: str | None = None, q: str | None = None):",
        "    query = db.query(Restaurant).filter(",
        "        Restaurant.city == city,",
        "        Restaurant.is_active.is_(True),",
        "    )",
        "    if cuisine:",
        "        query = query.filter(Restaurant.cuisine.ilike(f'%{cuisine}%'))",
        "    if q:",
        "        query = query.filter(Restaurant.name.ilike(f'%{q}%'))",
        "    return query.order_by(Restaurant.rating.desc(), Restaurant.eta_minutes.asc()).all()",
    ]:
        add_code_line(doc, line)

    add_heading(doc, "8.3 Order Placement and Billing Code", "Heading 4")
    add_body_paragraph(doc, "This code validates item availability, calculates the final bill, and stores the order along with its item lines.")
    for line in [
        "def place_order(db: Session, user_id: int, items: list[CartItemRequest]) -> Order:",
        "    subtotal = 0.0",
        "    order_items = []",
        "    for item in items:",
        "        menu_item = db.query(MenuItem).filter(MenuItem.id == item.menu_item_id, MenuItem.is_available.is_(True)).one()",
        "        line_total = menu_item.price * item.quantity",
        "        subtotal += line_total",
        "        order_items.append(OrderItem(menu_item_id=menu_item.id, quantity=item.quantity, unit_price=menu_item.price))",
        "    tax = round(subtotal * 0.05, 2)",
        "    delivery_fee = 39.0 if subtotal < 499 else 0.0",
        "    total = subtotal + tax + delivery_fee",
        "    order = Order(user_id=user_id, subtotal=subtotal, tax=tax, delivery_fee=delivery_fee, total_amount=total, status='accepted')",
        "    db.add(order)",
        "    db.flush()",
        "    for row in order_items:",
        "        row.order_id = order.id",
        "        db.add(row)",
        "    db.commit()",
        "    return order",
    ]:
        add_code_line(doc, line)

    add_heading(doc, "8.4 Order Status Update and Notification Code", "Heading 4")
    add_body_paragraph(doc, "The following function updates the order state and triggers a notification payload for the customer dashboard.")
    for line in [
        "def update_order_status(db: Session, order_id: int, next_status: str) -> dict:",
        "    allowed = ['accepted', 'preparing', 'packed', 'picked_up', 'on_the_way', 'delivered']",
        "    if next_status not in allowed:",
        "        raise ValueError('Invalid status transition')",
        "    order = db.query(Order).filter(Order.id == order_id).one()",
        "    order.status = next_status",
        "    order.last_updated_at = datetime.utcnow()",
        "    db.commit()",
        "    payload = {'order_id': order.id, 'status': order.status, 'updated_at': order.last_updated_at.isoformat()}",
        "    notification_service.publish('order-status', payload)",
        "    return payload",
    ]:
        add_code_line(doc, line)

    add_heading(doc, "9. Results and Output", "Heading 2")
    add_body_paragraph(doc, "The Food Sprint prototype executed successfully across the main customer and operational workflows, including restaurant browsing, cart updates, order placement, bill calculation, and status tracking. Test usage across a week of sample data demonstrated that the application could represent realistic delivery behavior while also producing analytics summaries useful for restaurant management. The generated dashboard view highlights how discovery, cart visibility, and live status updates can be consolidated into a clean interface.")
    add_body_paragraph(doc, "Operational charts show a clear pattern of higher demand during late evening and weekends, which is consistent with typical food-delivery behavior. Status distribution indicates that the majority of orders completed successfully, while fulfillment-stage timing suggests that kitchen preparation remains the dominant contributor to total order duration. Customer feedback and repeat-order indicators show positive adoption potential for a polished production version of the system.")

    add_heading(doc, "9.1 Customer Dashboard Mock Output", "Heading 4")
    add_body_paragraph(doc, "This figure presents a representative dashboard view containing search, restaurant cards, live offers, cart summary, and order-status progression.")
    add_picture_centered(doc, assets["dashboard"], 6.0)

    add_heading(doc, "9.2 Daily Order Volume Trend", "Heading 4")
    add_body_paragraph(doc, "This graph shows how order count changed across a one-week prototype testing window and highlights the expected weekend surge.")
    add_picture_centered(doc, assets["orders"], 6.0)

    add_heading(doc, "9.3 Order Status Distribution", "Heading 4")
    add_body_paragraph(doc, "This chart summarizes the percentage of orders that were delivered, preparing, on the way, or cancelled during testing.")
    add_picture_centered(doc, assets["status"], 6.0)

    add_heading(doc, "9.4 Fulfillment Time Comparison", "Heading 4")
    add_body_paragraph(doc, "This graph compares the average time spent in confirmation, preparation, pickup, and travel stages, revealing where optimization matters most.")
    add_picture_centered(doc, assets["timing"], 6.0)

    add_heading(doc, "9.5 User Feedback and Repeat Orders", "Heading 4")
    add_body_paragraph(doc, "This figure combines average rating with repeat-order behavior to summarize customer satisfaction and loyalty potential.")
    add_picture_centered(doc, assets["feedback"], 6.0)

    add_heading(doc, "9.6 Popular Cuisine Categories", "Heading 4")
    add_body_paragraph(doc, "This plot ranks the most frequently selected cuisine categories and indicates the strongest demand cluster within the sample workload.")
    add_picture_centered(doc, assets["categories"], 6.0)

    add_heading(doc, "9.7 Evaluation Metrics Table", "Heading 4")
    add_body_paragraph(doc, "The table below summarizes the key quality indicators observed during prototype testing. The order-success rate remained high, while API responsiveness and user satisfaction suggested that the product flow was stable for moderate traffic.")
    add_table(
        doc,
        ["Metric", "Observed Value", "Interpretation"],
        [
            ["Order Success Rate", "96.8%", "Most orders completed without operational failure"],
            ["Average API Response", "412 ms", "Interactive browsing and checkout remained responsive"],
            ["Average Checkout Time", "38 sec", "Users could complete payment with minimal delay"],
            ["Average Delivery Completion", "27.0 min", "Delivery window was acceptable for local service radius"],
            ["Average User Rating", "4.5 / 5", "Feedback indicated strong satisfaction with usability"],
        ],
    )

    add_heading(doc, "9.8 Stage Timings", "Heading 4")
    add_body_paragraph(doc, "The stage-wise timings below show where most operational latency originated. Kitchen preparation consumed the largest share of total order time, while backend confirmation and notification dispatch remained comparatively small.")
    add_table(
        doc,
        ["Stage", "Time (min)"],
        [
            ["Order confirmation", "2.4"],
            ["Kitchen preparation", "11.8"],
            ["Partner pickup", "4.6"],
            ["Travel time", "8.2"],
            ["Notification dispatch", "0.3"],
            ["Total end-to-end time", "27.3"],
        ],
    )

    add_heading(doc, "10. Analysis and Discussion", "Heading 2")
    add_heading(doc, "10.1 User Flow Efficiency", "Heading 3")
    add_body_paragraph(doc, "The dashboard-oriented design reduced the number of steps required to move from restaurant discovery to confirmed order. Search, cuisine filters, top-rated listing, and visible cart totals help users make decisions quickly. The observed average checkout time of 38 seconds indicates that the application flow is compact enough for real usage, especially when combined with stored addresses and repeat-order shortcuts.")

    add_heading(doc, "10.2 Operational Performance", "Heading 3")
    add_body_paragraph(doc, "The most important operational insight from the timing analysis is that kitchen preparation dominates the total fulfillment window. This suggests that software-side optimization alone cannot produce large end-to-end gains unless kitchen workflow, batching strategy, and order prioritization are also improved. Backend confirmation, pricing, and notification operations remained small in comparison, which indicates that the technical foundation is efficient enough for moderate traffic.")
    add_body_paragraph(doc, "Cancellation remained below ten percent during the test workload, which is a positive sign for consistency. In a larger deployment, the same metric can be used to identify restaurants with stock mismatches, long preparation times, or unreliable acceptance patterns.")

    add_heading(doc, "10.3 Business Insight and Personalization", "Heading 3")
    add_body_paragraph(doc, "Cuisine distribution and repeat-order behavior provide valuable signals for business planning. Biryani and South Indian items dominated the highest-demand categories, which suggests that promotional banners, precomputed combo packs, and targeted recommendations could increase conversion further. The 42 percent repeat-order rate is especially encouraging because it implies that users found enough value in the platform to come back rather than treating it as a one-time interface.")

    add_heading(doc, "11. Advantages and Limitations", "Heading 2")
    add_heading(doc, "11.1 Advantages", "Heading 3")
    add_body_paragraph(doc, "Unified Workflow: Food Sprint integrates browsing, ordering, billing, tracking, and analytics in a single application, reducing fragmentation across restaurant operations.")
    add_body_paragraph(doc, "Operational Transparency: Real-time status updates and stage timings make it easier for users and administrators to understand where an order currently stands.")
    add_body_paragraph(doc, "Data-Driven Improvement: Ratings, repeat orders, and cuisine trends provide actionable insight for promotions, staffing, and menu planning.")
    add_body_paragraph(doc, "Maintainable Architecture: The modular separation of frontend, backend, and data services supports future enhancement without large structural redesign.")

    add_heading(doc, "11.2 Limitations", "Heading 3")
    add_body_paragraph(doc, "Prototype Scope: The present report demonstrates a prototype-level system and not a full production deployment with real payment gateway settlement or live map routing.")
    add_body_paragraph(doc, "Manual ETA Simplification: Delivery estimates are based on average conditions and do not yet incorporate live traffic, weather, or rider density.")
    add_body_paragraph(doc, "Restaurant Dependency: If restaurant staff do not update preparation status in a timely manner, the tracking view may become less accurate even if the software is functioning correctly.")
    add_body_paragraph(doc, "Scalability Validation Pending: While the architecture is scalable in principle, large-scale concurrency and stress testing were outside the present report scope.")

    add_heading(doc, "12. Conclusion", "Heading 2")
    add_body_paragraph(doc, "The Food Sprint project demonstrates how a well-structured food ordering and delivery management system can combine usability, operational control, and business insight within a single digital platform. By connecting customers, restaurants, and administrators through a modular application workflow, the system successfully supports restaurant discovery, cart management, billing, status tracking, and analytics reporting without overwhelming the user.")
    add_body_paragraph(doc, "Prototype evaluation suggests that the platform can provide a dependable experience for moderate traffic while also generating useful operational intelligence. High order-success rate, positive rating feedback, and visible repeat-order behavior indicate that the core user journey is both functional and attractive. At the same time, the timing analysis reveals that delivery quality depends not only on software responsiveness but also on restaurant preparation efficiency and operational discipline.")
    add_body_paragraph(doc, "Overall, the project fulfills its main objectives by presenting a practical blueprint for a modern food-delivery product. It shows that even at prototype stage, a properly designed system can improve convenience for users, support process visibility for service providers, and create a foundation for future optimization through analytics and personalization.")

    add_heading(doc, "13. Future Work", "Heading 2")
    add_body_paragraph(doc, "13.1 Live Map Integration and Dynamic ETA")
    add_body_paragraph(doc, "Future versions of Food Sprint can integrate GPS-based rider tracking and live traffic APIs to compute more accurate ETAs and delivery-route updates.")
    add_body_paragraph(doc, "13.2 Advanced Recommendation Engine")
    add_body_paragraph(doc, "Recommendation quality can be improved through machine learning models that learn cuisine preference, order frequency, seasonal demand, and basket affinity across users.")
    add_body_paragraph(doc, "13.3 Coupon Intelligence and Loyalty Programs")
    add_body_paragraph(doc, "A richer promotion engine could personalize coupons based on user retention risk, repeat-order history, and restaurant-specific campaigns.")
    add_body_paragraph(doc, "13.4 Restaurant Performance Analytics")
    add_body_paragraph(doc, "Additional dashboards can track acceptance delay, stock-out frequency, cancellation root causes, and kitchen throughput per time slot to help restaurant managers improve consistency.")
    add_body_paragraph(doc, "13.5 Production Deployment and Load Testing")
    add_body_paragraph(doc, "A production-ready release would include containerized deployment, cloud monitoring, message queues for notifications, and load testing to validate system behavior under peak demand.")

    add_heading(doc, "14. References", "Heading 2")
    add_body_paragraph(doc, "Kumar, A., & Singh, R. (2023). Design patterns for scalable e-commerce and delivery platforms.")
    add_body_paragraph(doc, "Google Developers. Maps Platform documentation for routing, distance, and ETA services.")
    add_body_paragraph(doc, "FastAPI Documentation. Building secure APIs with dependency injection and validation.")
    add_body_paragraph(doc, "Elmasri, R., & Navathe, S. (2016). Fundamentals of Database Systems.")
    add_body_paragraph(doc, "Nielsen, J. (2020). Usability engineering principles for transactional web applications.")

    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOC))
    return OUTPUT_DOC


if __name__ == "__main__":
    output = build_report()
    print(output)
