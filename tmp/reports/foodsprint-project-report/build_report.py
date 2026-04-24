from __future__ import annotations

import ast
import math
import shutil
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\spand\OneDrive\Documents\New project")
OUT_DIR = ROOT / "outputs" / "foodsprint-project-report"
ASSET_DIR = OUT_DIR / "assets"
REPORT_PATH = OUT_DIR / "FoodSprint_Project_Report.docx"

FRONTEND_SRC = ROOT / "tmp" / "slides" / "foodsprint-evaluation-ppt" / "assets" / "frontend_home.png"

TITLE = "FoodSprint: AI-Assisted Food Ordering and Restaurant Discovery Web Application"
SUBTITLE = "A Project Report"


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
        "arialbd.ttf" if bold else "arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def parse_seed_constants() -> tuple[list[dict], list[tuple]]:
    source_path = ROOT / "app" / "services" / "data_seed.py"
    module = ast.parse(source_path.read_text(encoding="utf-8"))
    restaurants = []
    base_menu = []
    for node in module.body:
        if not isinstance(node, ast.Assign) or len(node.targets) != 1:
            continue
        target = node.targets[0]
        if not isinstance(target, ast.Name):
            continue
        if target.id == "RESTAURANTS":
            restaurants = ast.literal_eval(node.value)
        elif target.id == "BASE_MENU":
            base_menu = ast.literal_eval(node.value)
    return restaurants, base_menu


def extract_named_block(relative_path: str, name: str) -> str:
    path = ROOT / relative_path
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    tree = ast.parse(text)
    for node in ast.walk(tree):
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name == name:
            return "\n".join(lines[node.lineno - 1 : node.end_lineno])
    raise ValueError(f"Unable to find block {name!r} in {relative_path}")


def save_cropped_ui_images() -> dict[str, Path]:
    images = {}
    with Image.open(FRONTEND_SRC) as img:
        full = ASSET_DIR / "frontend_full.png"
        img.save(full)
        images["full"] = full

        discovery = img.crop((40, 0, img.width - 40, 2050))
        discovery_path = ASSET_DIR / "frontend_discovery.png"
        discovery.save(discovery_path)
        images["discovery"] = discovery_path

        catalog = img.crop((40, 1900, img.width - 40, 4050))
        catalog_path = ASSET_DIR / "frontend_catalog.png"
        catalog.save(catalog_path)
        images["catalog"] = catalog_path

    return images


def draw_bar_chart(
    data: dict[str, float],
    title: str,
    subtitle: str,
    out_path: Path,
    accent: str = "#FC8019",
    value_suffix: str = "",
) -> None:
    width, height = 1400, 820
    margin_left, margin_right = 130, 80
    margin_top, margin_bottom = 130, 120
    chart_w = width - margin_left - margin_right
    chart_h = height - margin_top - margin_bottom

    canvas = Image.new("RGB", (width, height), "#FFFDF8")
    draw = ImageDraw.Draw(canvas)

    title_font = get_font(40, bold=True)
    subtitle_font = get_font(22)
    label_font = get_font(20)
    value_font = get_font(22, bold=True)

    draw.text((margin_left, 38), title, font=title_font, fill="#18181B")
    draw.text((margin_left, 84), subtitle, font=subtitle_font, fill="#52525B")

    draw.line(
        [(margin_left, margin_top + chart_h), (margin_left + chart_w, margin_top + chart_h)],
        fill="#D4D4D8",
        width=3,
    )
    draw.line(
        [(margin_left, margin_top), (margin_left, margin_top + chart_h)],
        fill="#D4D4D8",
        width=3,
    )

    max_value = max(data.values()) if data else 1
    bar_count = max(len(data), 1)
    slot = chart_w / bar_count
    bar_w = min(120, slot * 0.55)

    for idx, (label, value) in enumerate(data.items()):
        x = margin_left + slot * idx + (slot - bar_w) / 2
        bar_h = 0 if max_value == 0 else chart_h * (value / max_value)
        y = margin_top + chart_h - bar_h
        draw.rounded_rectangle((x, y, x + bar_w, margin_top + chart_h), radius=20, fill=accent)

        value_text = f"{value:g}{value_suffix}"
        value_bbox = draw.textbbox((0, 0), value_text, font=value_font)
        draw.text(
            (x + (bar_w - (value_bbox[2] - value_bbox[0])) / 2, y - 34),
            value_text,
            font=value_font,
            fill="#18181B",
        )

        label_bbox = draw.multiline_textbbox((0, 0), label, font=label_font, spacing=4, align="center")
        draw.multiline_text(
            (
                x + (bar_w - (label_bbox[2] - label_bbox[0])) / 2,
                margin_top + chart_h + 20,
            ),
            label,
            font=label_font,
            fill="#3F3F46",
            spacing=4,
            align="center",
        )

    canvas.save(out_path)


def build_chart_assets(restaurants: list[dict], base_menu: list[tuple]) -> dict[str, Path]:
    city_counts = Counter(item["city"] for item in restaurants)
    food_type_counts = Counter(item[3 if False else 3] for item in base_menu)
    category_counts = Counter(item[2] for item in base_menu)

    city_chart = ASSET_DIR / "restaurants_by_city.png"
    draw_bar_chart(
        dict(city_counts),
        "Seeded Restaurants by City",
        "Counts derived from app/services/data_seed.py",
        city_chart,
        accent="#FC8019",
    )

    food_mix_chart = ASSET_DIR / "menu_food_type.png"
    scaled_food_type_counts = {key.title(): value * len(restaurants) for key, value in food_type_counts.items()}
    draw_bar_chart(
        scaled_food_type_counts,
        "Menu Mix by Food Type",
        "Each of the 11 restaurants receives the same 10-item base menu in the seed script",
        food_mix_chart,
        accent="#22C55E",
    )

    category_chart = ASSET_DIR / "menu_category_mix.png"
    scaled_category_counts = {key.title(): value * len(restaurants) for key, value in category_counts.items()}
    draw_bar_chart(
        scaled_category_counts,
        "Menu Mix by Category",
        "Protein, diet, balanced, and fast food tags support recommendation rules",
        category_chart,
        accent="#0EA5E9",
    )

    return {
        "cities": city_chart,
        "food_type": food_mix_chart,
        "category": category_chart,
    }


def set_page_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_default_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_centered_paragraph(document: Document, text: str, size: int = 12, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.style.font.name = "Times New Roman"


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.style.font.name = "Times New Roman"
    paragraph.paragraph_format.space_after = Pt(6)


def add_code_block(document: Document, title: str, code_text: str) -> None:
    add_paragraph(document, title)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F4F5")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "EDE9FE")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def add_picture_with_caption(document: Document, image_path: Path, caption: str, width: float = 6.0) -> None:
    document.add_picture(str(image_path), width=Inches(width))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def build_document(
    ui_images: dict[str, Path],
    charts: dict[str, Path],
    restaurants: list[dict],
    base_menu: list[tuple],
) -> None:
    city_counts = Counter(item["city"] for item in restaurants)
    total_menu_items = len(restaurants) * len(base_menu)
    food_type_counts = Counter(item[3] for item in base_menu)
    category_counts = Counter(item[2] for item in base_menu)

    create_app_code = extract_named_block("app/__init__.py", "create_app")
    distance_code = extract_named_block("app/services/location.py", "calculate_distance")
    nearby_code = extract_named_block("app/services/overpass.py", "fetch_overpass_restaurants")
    payment_code = extract_named_block("app/services/payments.py", "build_upi_uri")
    qr_code = extract_named_block("app/services/payments.py", "generate_qr_code_data_uri")
    recommendation_code = extract_named_block("app/services/recommendations.py", "build_diet_recommendation")

    document = Document()
    set_page_margins(document)
    set_default_styles(document)

    add_centered_paragraph(document, TITLE, size=18, bold=True)
    document.add_paragraph()
    add_centered_paragraph(document, SUBTITLE, size=15, bold=True)
    add_centered_paragraph(document, "Submitted in partial fulfillment of the course", size=12)
    document.add_paragraph()
    add_centered_paragraph(document, "Course Name: <COURSE NAME>", size=12)
    add_centered_paragraph(document, "Course Code: <COURSE CODE>", size=12)
    document.add_paragraph()
    add_centered_paragraph(document, "Submitted by:", size=12, bold=True)
    add_centered_paragraph(document, "<YOUR NAME>", size=12)
    add_centered_paragraph(document, "<ROLL NUMBER>", size=12)
    document.add_paragraph()
    add_centered_paragraph(document, "Under the supervision of", size=12, bold=True)
    add_centered_paragraph(document, "<GUIDE NAME>", size=12)
    add_centered_paragraph(document, "<DESIGNATION>", size=12)
    document.add_paragraph()
    add_centered_paragraph(document, "Academic Year: <YEAR>", size=12)
    document.add_paragraph()
    add_centered_paragraph(document, "Department of <DEPARTMENT>", size=12)
    add_centered_paragraph(document, "<UNIVERSITY NAME>", size=12)
    add_centered_paragraph(document, "<LOCATION>", size=12)

    document.add_page_break()

    add_heading(document, "INDEX", level=1)
    index_lines = [
        "1. Abstract",
        "2. Introduction",
        "3. Background and Concepts",
        "4. Problem Statement",
        "5. Objectives",
        "6. System Architecture",
        "7. Methodology",
        "8. Implementation with Code",
        "9. Results and Output",
        "10. Analysis and Discussion",
        "11. Advantages and Limitations",
        "12. Conclusion",
        "13. Future Work",
        "14. References",
    ]
    for line in index_lines:
        add_paragraph(document, line)

    document.add_page_break()

    add_heading(document, "1. ABSTRACT", level=1)
    add_paragraph(
        document,
        "FoodSprint is a Flask-based food ordering and restaurant discovery web application focused on Telangana users. "
        "The project addresses the problem of finding nearby restaurants, browsing menus, placing orders, completing "
        "a UPI-based payment flow, and tracking order progress in a single lightweight system. The implementation uses "
        "Flask, SQLAlchemy, Jinja templates, JavaScript, Leaflet maps, OpenStreetMap tiles, the Overpass API, and QR "
        "code generation. The application dataset is a seeded internal catalog containing 11 restaurants and 110 menu "
        "items distributed across Hyderabad, Warangal, Karimnagar, Nizamabad, and Khammam, and it is supplemented by "
        "live nearby places from OpenStreetMap. Local validation showed successful responses for the homepage, search "
        "API, and diet recommendation page, with a measured startup time of 1.961 seconds on the current machine. "
        "Instead of using ML model training, the project applies deterministic recommendation rules, caching, and "
        "client-side filtering for fast response and simple maintenance."
    )

    add_heading(document, "2. INTRODUCTION", level=1)
    add_paragraph(
        document,
        "Food delivery and restaurant discovery platforms are highly relevant because users increasingly expect fast "
        "search, location-aware results, digital payments, and order visibility from a single interface. In many "
        "student projects, these features are implemented in isolation, which leaves the system incomplete. FoodSprint "
        "was designed to connect them into one end-to-end workflow: search nearby restaurants, view menus, manage a "
        "cart, pay through UPI QR, and track status until delivery. The problem is challenging because nearby discovery "
        "depends on geographic data, ordering requires session-aware state handling, and operational dashboards must "
        "work for admin and staff roles. The chosen approach is suitable because Flask offers rapid backend development, "
        "SQLAlchemy provides a structured data model, and OpenStreetMap/Overpass provide accessible location data for "
        "map-driven restaurant search."
    )

    add_heading(document, "2.1 Literature Review", level=2)
    add_paragraph(
        document,
        "Existing platforms such as Swiggy and Zomato have shown that restaurant discovery must be fast, location-aware, "
        "and mobile-friendly. On the technical side, Flask is widely used for educational full-stack applications due "
        "to its simplicity and extension ecosystem, while Leaflet and OpenStreetMap are common choices for lightweight "
        "map interfaces. Rule-based recommendation systems are also frequently used in early-stage products because they "
        "are explainable, fast, and do not require large training datasets. FoodSprint combines these ideas into a "
        "student-friendly system: open map discovery, curated catalog data, session cart management, a simple UPI QR "
        "payment experience, and rule-based nutrition-aware suggestions."
    )

    add_heading(document, "3. BACKGROUND AND CONCEPTS", level=1)
    add_heading(document, "3.1 Flask-Based Web Application Architecture", level=2)
    add_paragraph(
        document,
        "The project uses Flask as the web framework, with route functions acting as controllers, Jinja templates "
        "rendering HTML views, and SQLAlchemy modeling database entities. This layered design makes the application "
        "easy to understand and extend."
    )
    add_heading(document, "3.2 Geospatial Restaurant Discovery", level=2)
    add_paragraph(
        document,
        "FoodSprint uses browser geolocation, city fallback coordinates, the Haversine distance formula, Leaflet maps, "
        "and the Overpass API to show nearby restaurants. This allows the app to combine fixed catalog data with live "
        "map-based context."
    )
    add_heading(document, "3.3 Session Cart and Order Lifecycle", level=2)
    add_paragraph(
        document,
        "Customer actions are stored in the session cart until checkout. After confirmation, the cart is transformed "
        "into an order, payment records are attached, and staff can move the order through statuses such as PLACED, "
        "ACCEPTED, PREPARING, OUT_FOR_DELIVERY, and DELIVERED."
    )
    add_heading(document, "3.4 Rule-Based Recommendation Engine", level=2)
    add_paragraph(
        document,
        "Instead of training a predictive model, FoodSprint uses deterministic rules for diet-aware suggestions. "
        "Menu items are tagged using fields such as category, calories, food type, and healthy badge, and then "
        "ranked according to user goals such as weight loss, muscle gain, balanced diet, or vegetarian preference."
    )
    add_heading(document, "3.5 QR-Based UPI Payment Integration", level=2)
    add_paragraph(
        document,
        "The payment module constructs a UPI URI, converts it into an SVG QR code, and presents it during checkout. "
        "This keeps the payment flow simple and demo-friendly while matching a real Indian payment pattern."
    )

    add_heading(document, "4. PROBLEM STATEMENT", level=1)
    add_paragraph(
        document,
        "The problem addressed by FoodSprint is the lack of a compact, complete, and locally relevant restaurant "
        "discovery and ordering system for student demonstration purposes. Many small projects either stop at static "
        "restaurant listing pages or provide only basic cart functionality without maps, payment, or operations support. "
        "The problem is difficult because the system must combine different concerns: geographic lookup, menu browsing, "
        "session handling, order creation, review management, and role-based dashboards. The system must therefore "
        "provide accurate restaurant discovery, a smooth ordering workflow, accessible payment support, and clear "
        "administrative control."
    )

    add_heading(document, "5. OBJECTIVES", level=1)
    objectives = [
        "Provide location-aware restaurant discovery using both seeded Telangana data and nearby OpenStreetMap places.",
        "Support a complete customer workflow from account creation to cart management, checkout, payment, and order tracking.",
        "Offer diet-aware and repeat-order recommendations using transparent rule-based logic.",
        "Enable admin and staff roles to manage UPI settings, staff accounts, and order status updates.",
    ]
    for item in objectives:
        add_bullet(document, item)

    add_heading(document, "6. SYSTEM ARCHITECTURE", level=1)
    add_paragraph(
        document,
        "The system architecture combines frontend rendering, backend routing, service modules, a relational database, "
        "and external APIs. Customer requests are received through HTML pages or JSON endpoints. The backend routes call "
        "helper services for distance calculation, nearby place lookup, recommendations, and payment QR generation. "
        "Persistent entities such as users, restaurants, menu items, orders, payments, and reviews are stored using "
        "SQLAlchemy models. The final outputs are rendered pages, interactive map results, order records, and admin/staff dashboards."
    )
    add_code_block(
        document,
        "High-level pipeline:",
        "User Browser\n"
        "    -> Flask Routes\n"
        "        -> Services (Location, Overpass, Recommendations, Payments)\n"
        "            -> SQLAlchemy Models / SQLite or PostgreSQL\n"
        "                -> HTML Pages + JSON Responses + UPI QR + Order Status Updates"
    )

    add_heading(document, "7. METHODOLOGY", level=1)
    add_heading(document, "7.1 Dataset Loading", level=2)
    add_paragraph(
        document,
        f"The project does not rely on an external ML dataset. Instead, it initializes a structured internal dataset in "
        f"`app/services/data_seed.py` containing {len(restaurants)} restaurants and a reusable base menu of {len(base_menu)} "
        f"items per restaurant, resulting in {total_menu_items} menu items after seeding. The seed script also creates one "
        f"admin account and one staff account for demonstration."
    )
    add_heading(document, "7.2 Data Cleaning", level=2)
    add_paragraph(
        document,
        "Input values such as city names, keywords, and food type filters are normalized using lowercase string handling. "
        "Nearby place results from Overpass are deduplicated by name and rounded coordinates, and entries missing latitude "
        "or longitude are discarded. When browser geolocation is unavailable, the app falls back to predefined Telangana city coordinates."
    )
    add_heading(document, "7.3 Feature Engineering", level=2)
    add_paragraph(
        document,
        "The application derives useful features such as distance_km, rating badges, healthy tags, calories, cuisine "
        "labels, and veg/non-veg categories. These fields support filtering, ranking, and recommendation logic across the UI."
    )
    add_heading(document, "7.4 Data Transformation", level=2)
    add_paragraph(
        document,
        "Distances are computed using the Haversine formula, search results are serialized into JSON-friendly restaurant "
        "cards, and payment details are converted into a UPI URI and then an SVG QR data URI. This transformation layer "
        "bridges database records and user-facing interfaces."
    )
    add_heading(document, "7.5 Model Design", level=2)
    add_paragraph(
        document,
        "FoodSprint uses a rule-based recommendation engine instead of a trained ML model. Diet rules specify which "
        "categories to recommend or avoid for goals such as weight loss or muscle gain, while history-based suggestions "
        "use keyword counting over recent order items. This model is deterministic, fast, and explainable."
    )
    add_heading(document, "7.6 Training Strategy", level=2)
    add_paragraph(
        document,
        "A supervised training pipeline is not applicable in this project because the recommendation system is heuristic "
        "rather than learned from labeled examples. Validation was performed through local route testing, seed-data checks, "
        "and manual inspection of the rendered pages and user flows."
    )
    add_heading(document, "7.7 Evaluation Metrics", level=2)
    add_paragraph(
        document,
        "Evaluation focused on functional completeness and route behavior rather than predictive accuracy. The report "
        "records dataset coverage, HTTP status codes for important pages, seeded entity counts, and response times for "
        "representative routes such as the homepage, search API, and diet recommendation page."
    )
    add_heading(document, "7.8 Optimization / HPC", level=2)
    add_paragraph(
        document,
        "No specialized GPU or HPC infrastructure is required for this application. Performance optimizations are "
        "lightweight but useful: the distance helper uses an LRU cache, the Overpass service caches results for 300 "
        "seconds, nearby markers are grouped with Leaflet marker clustering, and client-side filtering reduces repeat requests."
    )

    add_heading(document, "8. IMPLEMENTATION WITH CODE", level=1)
    add_heading(document, "8.1 Application Bootstrapping Code", level=2)
    add_code_block(document, "Code extracted from app/__init__.py", create_app_code)
    add_heading(document, "8.2 Nearby Discovery and Distance Calculation", level=2)
    add_code_block(document, "Distance helper from app/services/location.py", distance_code)
    add_code_block(document, "Nearby place fetch from app/services/overpass.py", nearby_code)
    add_heading(document, "8.3 UPI Payment and QR Generation", level=2)
    add_code_block(document, "UPI payload builder from app/services/payments.py", payment_code)
    add_code_block(document, "QR generator from app/services/payments.py", qr_code)
    add_heading(document, "8.4 Recommendation Logic", level=2)
    add_code_block(document, "Diet recommendation logic from app/services/recommendations.py", recommendation_code)

    add_heading(document, "9. RESULTS AND OUTPUT", level=1)
    add_heading(document, "9.1 Output Screenshots", level=2)
    add_paragraph(
        document,
        "The following screenshots were taken from the running FoodSprint interface used during the PPT preparation stage. "
        "They show the homepage, nearby map discovery section, and the restaurant listing experience."
    )
    add_picture_with_caption(document, ui_images["full"], "Figure 1. Full FoodSprint homepage with hero, nearby discovery, and restaurant sections.", width=4.6)
    add_picture_with_caption(document, ui_images["discovery"], "Figure 2. Discovery section showing hero content, filters, and the map-based nearby restaurant panel.", width=6.1)
    add_picture_with_caption(document, ui_images["catalog"], "Figure 3. Restaurant listing sections showing top-rated and statewide ordering cards.", width=6.1)

    add_heading(document, "9.2 Graphs", level=2)
    add_paragraph(
        document,
        "Because FoodSprint is a web application rather than a predictive ML model, charts such as Actual vs Predicted "
        "and Training Loss are not applicable. Instead, the following graphs summarize project data coverage and system-relevant distributions."
    )
    add_picture_with_caption(document, charts["cities"], "Figure 4. Distribution of seeded restaurants across Telangana cities.", width=6.2)
    add_picture_with_caption(document, charts["food_type"], "Figure 5. Distribution of seeded menu items by food type.", width=6.2)
    add_picture_with_caption(document, charts["category"], "Figure 6. Distribution of seeded menu items by recommendation category.", width=6.2)

    add_heading(document, "9.3 Metrics Table", level=2)
    add_table(
        document,
        ["Metric", "Value", "Observation"],
        [
            ["Seeded restaurants", str(len(restaurants)), "Initialized from the Telangana-focused seed script"],
            ["Seeded menu items", str(total_menu_items), "Generated as 10 items per restaurant"],
            ["Cities covered", str(len(city_counts)), "Hyderabad, Warangal, Karimnagar, Nizamabad, and Khammam"],
            ["Food type ratio", f"Veg {food_type_counts['veg']} : Non-Veg {food_type_counts['non-veg']}", "Base menu composition before multiplying by restaurant count"],
            ["Homepage status", "200", "Successful local test-client response"],
            ["Search API status", "200", "Successful local query for biryani"],
            ["Diet page status", "200", "Balanced diet page rendered successfully"],
        ],
    )

    add_heading(document, "9.4 Execution Time", level=2)
    add_paragraph(
        document,
        "The following timings were measured locally on the current machine using Python and Flask test-client execution. "
        "They are useful for relative comparison but may vary on other systems."
    )
    add_table(
        document,
        ["Stage", "Time", "Note"],
        [
            ["Application startup", "1.961 s", "Includes app creation, database init, and seed checks"],
            ["Homepage route (/)", "52.72 ms", "Rendered via Flask test client"],
            ["Search API (/api/search?q=biryani)", "2.12 ms", "JSON search response"],
            ["Diet page (/diet?goal=balanced_diet)", "25.40 ms", "Recommendation page rendering"],
        ],
    )

    add_heading(document, "10. ANALYSIS AND DISCUSSION", level=1)
    add_paragraph(
        document,
        "The hybrid nearby-discovery workflow is the strongest part of the project because it combines curated data and live map context. "
        "The seeded catalog guarantees that the application always has meaningful restaurants to display, while the Overpass integration "
        "adds real nearby places when network access is available. The rule-based diet recommendation engine also performs well for a student "
        "project because it is transparent and fast, but it does not yet offer the personalization depth of a trained recommendation model. "
        "Response times for the tested routes were comfortably low in local execution, especially for the search API. The main trade-off is "
        "that the system depends on external services and static image mappings for parts of the experience."
    )

    add_heading(document, "11. ADVANTAGES AND LIMITATIONS", level=1)
    add_heading(document, "11.1 Advantages", level=2)
    advantages = [
        "Covers a complete end-to-end workflow: discovery, cart, payment, tracking, and operations.",
        "Uses open map technologies, which makes nearby restaurant discovery realistic and visually strong.",
        "Seeds the database automatically, so the project is easy to demo without manual setup.",
        "Provides explainable recommendation logic using nutrition tags and order history keywords.",
        "Supports multiple actors: customer, admin, and staff.",
    ]
    for item in advantages:
        add_bullet(document, item)

    add_heading(document, "11.2 Limitations", level=2)
    limitations = [
        "The recommendation engine is rule-based and not learned from real user behavior.",
        "The live nearby-search experience depends on OpenStreetMap/Overpass availability.",
        "The default payment flow is a demo-friendly manual UPI confirmation rather than a production gateway.",
        "Some food and restaurant images rely on third-party URLs, so image quality may vary.",
        "Course, student, and guide details still need to be filled manually in the title page placeholders.",
    ]
    for item in limitations:
        add_bullet(document, item)

    add_heading(document, "12. CONCLUSION", level=1)
    add_paragraph(
        document,
        "FoodSprint demonstrates how a student software project can go beyond a static restaurant list and become a realistic food-ordering system. "
        "The application integrates restaurant discovery, order placement, UPI QR checkout, tracking, recommendations, and operational dashboards "
        "within a single Flask codebase. The project achieved its primary goal of providing a full-stack, presentation-ready web application that "
        "is easy to run locally and clear enough to explain in an academic setting. The main takeaway is that careful use of open APIs, seeded data, "
        "and modular Flask services can produce a strong end-to-end demo without requiring a large ML pipeline."
    )

    add_heading(document, "13. FUTURE WORK", level=1)
    future_work = [
        "Integrate a production-grade payment gateway such as Razorpay end-to-end instead of manual payment confirmation.",
        "Add delivery assignment, real-time socket updates, and push notifications for order tracking.",
        "Replace rule-based recommendations with a hybrid ML plus behavioral recommendation system when real user data is available.",
        "Improve image handling with locally stored assets or a managed CDN for higher reliability.",
        "Deploy the system publicly with PostgreSQL, CI validation, and monitoring dashboards.",
    ]
    for item in future_work:
        add_bullet(document, item)

    add_heading(document, "14. REFERENCES", level=1)
    references = [
        "Flask Documentation: https://flask.palletsprojects.com/",
        "SQLAlchemy Documentation: https://docs.sqlalchemy.org/",
        "Leaflet Documentation: https://leafletjs.com/",
        "OpenStreetMap: https://www.openstreetmap.org/",
        "Overpass API: https://overpass-api.de/",
        "qrcode Python Package: https://pypi.org/project/qrcode/",
        "Project repository files: README.md, app/routes.py, app/models.py, app/services/data_seed.py, app/services/location.py, app/services/overpass.py, app/services/payments.py, app/services/recommendations.py",
        "Captured FoodSprint UI screenshots and project PPT assets generated in the local workspace",
    ]
    for ref in references:
        add_bullet(document, ref)

    document.save(str(REPORT_PATH))


def main() -> None:
    ensure_dirs()
    restaurants, base_menu = parse_seed_constants()
    ui_images = save_cropped_ui_images()
    charts = build_chart_assets(restaurants, base_menu)
    build_document(ui_images, charts, restaurants, base_menu)
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
